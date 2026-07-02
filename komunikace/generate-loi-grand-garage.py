#!/usr/bin/env python3
"""Generátor Letter of Intent — Grand Garage Linz pro EHMK 2028."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.expanduser("~/Desktop/EHMK28/komunikace/LoI-Grand-Garage-draft.docx")

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY = RGBColor(0xF5, 0xF5, 0xF5)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), 'CCCCCC')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_no_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:color'), 'auto')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_cell_content(cell, label_text, value_text, bg="F8F8F8"):
    set_cell_bg(cell, bg)
    set_cell_borders(cell)
    cell.top_margin = Cm(0.3)
    cell.bottom_margin = Cm(0.3)
    cell.left_margin = Cm(0.3)
    cell.right_margin = Cm(0.3)
    # Label
    p_label = cell.paragraphs[0]
    run = p_label.add_run(label_text)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    # Value
    p_val = cell.add_paragraph()
    run2 = p_val.add_run(value_text)
    run2.font.size = Pt(10)
    p_val.paragraph_format.space_after = Pt(2)

doc = Document()

# Okraje stránky
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# Nadpis
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
run = title.add_run("Letter of Intent")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = NAVY

# Podtitul
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(20)
run = subtitle.add_run(
    "Intention to cooperate with an international partner on the project Make & Design Lab"
)
run.italic = True
run.font.size = Pt(11)
run.font.color.rgb = GRAY

# Tabulka žadatel / partner (2 sloupce)
table = doc.add_table(rows=0, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

def add_header_row(tbl, left, right):
    row = tbl.add_row()
    for i, text in enumerate([left, right]):
        cell = row.cells[i]
        set_cell_bg(cell, "1B2A4A")
        set_cell_borders(cell)
        cell.top_margin = Cm(0.2)
        cell.bottom_margin = Cm(0.2)
        cell.left_margin = Cm(0.3)
        cell.right_margin = Cm(0.3)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = WHITE

add_header_row(table, "Applicant", "Partner")

rows_data = [
    ("Name", "Make More s.r.o.", "Name", "Grand Garage (TGW Future Wings GmbH)"),
    ("Address", "[DOPLNIT — adresa provozovny Make More v ČB]", "Address", "Peter-Behrens-Platz 10\n4020 Linz, Austria"),
    ("Company ID", "[DOPLNIT — IČO Make More s.r.o.]", "Company ID", "[DOPLNIT — Austrian registration no.]"),
    ("Telephone", "[DOPLNIT — Ondřej Kašpárek]", "Telephone", "[DOPLNIT]"),
    ("Email", "ondrej@makemore.cz", "Email", "[DOPLNIT — Grand Garage kontakt]"),
]

for l_label, l_val, r_label, r_val in rows_data:
    row = table.add_row()
    add_cell_content(row.cells[0], l_label, l_val)
    add_cell_content(row.cells[1], r_label, r_val)

doc.add_paragraph()

# Povinný text z šablony
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
run = p.add_run(
    "The applicant and partner declare that they have a common interest in preparing, "
    "organising and implementing a cultural programme as part of the European Capital of "
    "Culture 2028 in České Budějovice."
)
run.font.size = Pt(10.5)

# Popis partnerství (z draftu emailu)
p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(10)
run2 = p2.add_run(
    "Grand Garage (TGW Future Wings GmbH), operating one of Austria's leading makerspaces "
    "in the historic Tabakfabrik Linz, will contribute to the project Make & Design Lab as "
    "a knowledge and mentorship partner. Grand Garage brings expertise in community-centred "
    "makerspace operations, hands-on educational programming for youth and schools, and the "
    "development of maker communities rooted in the collaboration between industry, education, "
    "and the public."
)
run2.font.size = Pt(10.5)

p3 = doc.add_paragraph()
p3.paragraph_format.space_after = Pt(6)
run3 = p3.add_run("The partnership is envisioned in two phases:")
run3.font.size = Pt(10.5)

p4 = doc.add_paragraph()
p4.paragraph_format.space_after = Pt(6)
run4a = p4.add_run("2027 — Preparation phase: ")
run4a.bold = True
run4a.font.size = Pt(10.5)
run4b = p4.add_run(
    "Grand Garage will share expertise and operational know-how to support the design of "
    "the educational programme, curriculum, and community model for Make & Design Lab. "
    "This may include consultation sessions, a study visit to Linz, and advisory input as "
    "the pop-up makerspace concept is developed for the South Bohemian context."
)
run4b.font.size = Pt(10.5)

p5 = doc.add_paragraph()
p5.paragraph_format.space_after = Pt(10)
run5a = p5.add_run("2028 — Programme year: ")
run5a.bold = True
run5a.font.size = Pt(10.5)
run5b = p5.add_run(
    "Grand Garage will participate in the programme as a guest contributor — through a "
    "workshop session, lecture, or panel discussion — bringing international makerspace "
    "practice and experience to the South Bohemian region and European Capital of Culture "
    "2028 audiences."
)
run5b.font.size = Pt(10.5)

p6 = doc.add_paragraph()
p6.paragraph_format.space_after = Pt(20)
run6 = p6.add_run(
    "Both parties see this collaboration as a meaningful exchange between two organisations "
    "committed to the conviction that access to making, technology and creative education "
    "is a public good that belongs at the heart of a community."
)
run6.font.size = Pt(10.5)

# Podpisový blok
sig_table = doc.add_table(rows=2, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Místo a datum
row0 = sig_table.rows[0]
for cell in row0.cells:
    set_no_borders(cell)
    cell.top_margin = Cm(0.1)
    cell.bottom_margin = Cm(0.1)

cell_l = row0.cells[0]
p = cell_l.paragraphs[0]
p.add_run("In České Budějovice").font.size = Pt(10)
p2 = cell_l.add_paragraph()
p2.add_run("On ______________________").font.size = Pt(10)

cell_r = row0.cells[1]
p = cell_r.paragraphs[0]
p.add_run("In Linz").font.size = Pt(10)
p2 = cell_r.add_paragraph()
p2.add_run("On ______________________").font.size = Pt(10)

# Podpisy
row1 = sig_table.rows[1]
for cell in row1.cells:
    set_no_borders(cell)
    cell.top_margin = Cm(1.5)

cell_l = row1.cells[0]
p = cell_l.paragraphs[0]
p.paragraph_format.space_before = Pt(36)
# Čára nad podpisem
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
top = OxmlElement('w:top')
top.set(qn('w:val'), 'single')
top.set(qn('w:sz'), '4')
top.set(qn('w:color'), '000000')
pBdr.append(top)
pPr.append(pBdr)
run = p.add_run("Signature of Applicant")
run.font.size = Pt(9)
run.font.color.rgb = GRAY
p2 = cell_l.add_paragraph()
p2.add_run("Ondřej Kašpárek, Make More s.r.o.").font.size = Pt(9)

cell_r = row1.cells[1]
p = cell_r.paragraphs[0]
p.paragraph_format.space_before = Pt(36)
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
top = OxmlElement('w:top')
top.set(qn('w:val'), 'single')
top.set(qn('w:sz'), '4')
top.set(qn('w:color'), '000000')
pBdr.append(top)
pPr.append(pBdr)
run = p.add_run("Signature of the Foreign Partner")
run.font.size = Pt(9)
run.font.color.rgb = GRAY
p2 = cell_r.add_paragraph()
p2.add_run("[Name, Title]\nGrand Garage / TGW Future Wings GmbH").font.size = Pt(9)

doc.save(OUTPUT)
print(f"Uloženo: {OUTPUT}")
